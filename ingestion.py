"""
Document ingestion module.
Extracts clean text from PDF, DOCX, and TXT files.
Each extractor returns a list of (text, metadata) tuples so we can
track provenance (source filename, page number) through the pipeline.
"""

import io
import re
from dataclasses import dataclass

import PyPDF2
import docx


@dataclass
class Document:
    """A single logical unit of extracted text with provenance metadata."""
    text: str
    metadata: dict  # e.g. {"source": "file.pdf", "page": 3}


def clean_text(text: str) -> str:
    """
    Normalize whitespace and remove control characters.
    Keeps the text human-readable while removing noise that
    would hurt embedding quality.
    """
    # Replace multiple whitespace (including newlines) with single space
    text = re.sub(r"\s+", " ", text)
    # Remove non-printable control chars (keep newlines and tabs for structure)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    return text.strip()


def extract_pdf(file_bytes: bytes, filename: str) -> list[Document]:
    """Extract text page-by-page from a PDF so we preserve page-level metadata."""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    documents = []
    for page_num, page in enumerate(reader.pages, start=1):
        raw = page.extract_text() or ""
        text = clean_text(raw)
        if text:  # skip blank pages
            documents.append(Document(
                text=text,
                metadata={"source": filename, "page": page_num},
            ))
    return documents


def extract_docx(file_bytes: bytes, filename: str) -> list[Document]:
    """
    Extract text from DOCX. Word documents don't have a strong
    page concept, so we treat each paragraph as a unit and group
    non-empty paragraphs into a single document.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    text = clean_text(full_text)
    if not text:
        return []
    return [Document(text=text, metadata={"source": filename, "page": 1})]


def extract_txt(file_bytes: bytes, filename: str) -> list[Document]:
    """Plain text — decode and clean."""
    # Try UTF-8 first, fall back to latin-1 which never fails
    try:
        raw = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        raw = file_bytes.decode("latin-1")
    text = clean_text(raw)
    if not text:
        return []
    return [Document(text=text, metadata={"source": filename, "page": 1})]


# Dispatcher keyed by file extension
EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_txt,
}


def ingest_file(file_bytes: bytes, filename: str) -> list[Document]:
    """
    Main entry point: given raw file bytes and a filename,
    detect the type and extract documents.
    """
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(EXTRACTORS.keys())}")
    return extractor(file_bytes, filename)
